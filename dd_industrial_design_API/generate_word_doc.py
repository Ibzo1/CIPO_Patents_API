#!/usr/bin/env python3
"""
Generate Word Document from Technical Documentation
This script creates a formatted Word document for knowledge transfer.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def add_heading(doc, text, level=1):
    """Add a heading with proper formatting"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, style=None):
    """Add a paragraph with optional style"""
    if style:
        return doc.add_paragraph(text, style=style)
    return doc.add_paragraph(text)

def add_code_block(doc, code):
    """Add a formatted code block"""
    para = doc.add_paragraph()
    para.style = 'Code'
    run = para.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    # Light gray background effect via shading would require deeper XML manipulation
    return para

def add_table_from_markdown(doc, lines):
    """Parse markdown table and create Word table"""
    # Remove separator line
    table_lines = [line for line in lines if not re.match(r'\|[\s\-:]+\|', line)]

    if len(table_lines) < 2:
        return None

    # Parse header
    headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
    num_cols = len(headers)

    # Parse rows
    rows = []
    for line in table_lines[1:]:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if len(cells) == num_cols:
            rows.append(cells)

    if not rows:
        return None

    # Create table
    table = doc.add_table(rows=len(rows) + 1, cols=num_cols)
    table.style = 'Light Grid Accent 1'

    # Add headers
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        # Bold header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Add data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_data in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = cell_data

    return table

def setup_styles(doc):
    """Setup custom styles for the document"""
    styles = doc.styles

    # Code style
    try:
        code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_font = code_style.font
        code_font.name = 'Courier New'
        code_font.size = Pt(9)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
        code_style.paragraph_format.left_indent = Inches(0.5)
    except:
        # Style already exists
        pass

def generate_word_document():
    """Main function to generate the Word document"""

    # Create document
    doc = Document()

    # Setup styles
    setup_styles(doc)

    # Title page
    title = doc.add_heading('Industrial Design API', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Technical Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].font.color.rgb = RGBColor(68, 68, 68)

    doc.add_paragraph()

    date_para = doc.add_paragraph('Document Date: 2025-11-25')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    version_para = doc.add_paragraph('Version 1.0')
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Table of Contents placeholder
    doc.add_heading('Table of Contents', 1)
    doc.add_paragraph('(Generated automatically when opened in Microsoft Word)')
    doc.add_paragraph('In Microsoft Word: References → Table of Contents → Automatic Table')
    doc.add_page_break()

    # 1. Executive Summary
    add_heading(doc, '1. Executive Summary', 1)
    add_paragraph(doc,
        'The Industrial Design API is a Django REST Framework-based web service that provides '
        'read-only access to Canadian Industrial Design registration data from the Canadian '
        'Intellectual Property Office (CIPO). The API exposes comprehensive industrial design '
        'application and assignment information through RESTful endpoints with built-in filtering, '
        'search, and pagination capabilities.')

    add_heading(doc, 'Key Technologies:', 2)
    items = [
        'Django 5.1.1',
        'Django REST Framework 3.14.0',
        'PostgreSQL Database',
        'drf-yasg (Swagger/OpenAPI documentation)',
        'Waitress WSGI Server'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # 2. Project Purpose
    add_heading(doc, '2. Project Purpose', 1)
    add_paragraph(doc,
        'This API serves as a data access layer for industrial design information, enabling:')

    purposes = [
        'Programmatic access to industrial design applications and registrations',
        'Search and filter capabilities across multiple data dimensions',
        'Integration with external systems requiring industrial design data',
        'Support for applications needing Canadian IP data'
    ]
    for purpose in purposes:
        doc.add_paragraph(purpose, style='List Bullet')

    add_paragraph(doc,
        'The system is designed as a read-only API, ensuring data integrity while providing '
        'flexible query capabilities.')

    # 3. System Architecture
    add_heading(doc, '3. System Architecture', 1)

    add_heading(doc, '3.1 Application Structure', 2)
    structure = """dd_industrial_design_API/
├── DB_Main/                    # Django project configuration
│   ├── settings.py            # Application settings
│   ├── urls.py                # Root URL configuration
│   ├── wsgi.py                # WSGI configuration
│   └── asgi.py                # ASGI configuration
├── industrial_designs/         # Main application
│   ├── models.py              # Database models
│   ├── views.py               # API viewsets
│   ├── serializers.py         # Data serializers
│   ├── urls.py                # URL routing
│   └── migrations/            # Database migrations
├── serve.py                   # Production server entry point
├── manage.py                  # Django management script
└── requirements.txt           # Python dependencies"""

    add_code_block(doc, structure)

    add_heading(doc, '3.2 Technology Stack', 2)
    table_data = [
        '| Component | Technology | Version | Purpose |',
        '|-----------|-----------|---------|---------|',
        '| Web Framework | Django | 5.1.1 | Core application framework |',
        '| API Framework | Django REST Framework | 3.14.0 | RESTful API implementation |',
        '| Database | PostgreSQL | - | Primary data storage |',
        '| Database Adapter | psycopg2-binary | 2.9.10 | PostgreSQL connectivity |',
        '| API Documentation | drf-yasg | 1.21.7 | Swagger/OpenAPI docs |',
        '| Filtering | django-filter | 23.1 | Advanced query filtering |',
        '| Production Server | Waitress | - | WSGI HTTP server |',
        '| CORS Support | django-cors-headers | 4.6.0 | Cross-origin requests |'
    ]
    add_table_from_markdown(doc, table_data)

    add_heading(doc, '3.3 Database Configuration', 2)
    add_paragraph(doc, 'Database: IndustrialDesign')
    add_paragraph(doc, 'Schema: id_csv_2024_03_07')
    add_paragraph(doc, 'Connection Details:')
    db_details = ['Host: localhost', 'Port: 5432', 'User: Azhari']
    for detail in db_details:
        doc.add_paragraph(detail, style='List Bullet')

    add_paragraph(doc,
        'All models utilize the unmanaged model pattern (managed = False), meaning Django does '
        'not control database schema creation or migrations. The database schema is maintained externally.')

    # 4. Data Models
    add_heading(doc, '4. Data Models', 1)
    add_paragraph(doc,
        'The system contains 10 primary models organized into two categories:')

    add_heading(doc, '4.1 Application Models', 2)

    add_heading(doc, 'ApplicationMain', 3)
    add_paragraph(doc,
        'The primary model containing core industrial design application information.')
    add_paragraph(doc, 'Key Fields:')
    fields = [
        'application_number - Unique application identifier',
        'design_title - Title of the design',
        'design_status - Current status of the application',
        'filing_date - Date application was filed',
        'publication_date - Date design was published',
        'registration_date - Date design was registered',
        'registration_number - Assigned registration number',
        'registration_expiry_date - When registration expires',
        'priority_date - Priority claim date',
        'international_registration_number - Hague system number'
    ]
    for field in fields:
        doc.add_paragraph(field, style='List Bullet')

    add_heading(doc, 'ApplicationClassification', 3)
    add_paragraph(doc, 'Product classification information for applications.')
    add_paragraph(doc, 'Key Fields:')
    fields = [
        'application_number - Links to ApplicationMain',
        'classification_number - International classification code',
        'classification_kind - Type of classification',
        'product_description - Description of the product'
    ]
    for field in fields:
        doc.add_paragraph(field, style='List Bullet')

    add_heading(doc, 'ApplicationInterestedParty', 3)
    add_paragraph(doc,
        'Parties involved in the application (owners, designers, agents).')
    add_paragraph(doc, 'Key Fields:')
    fields = [
        'application_number - Links to ApplicationMain',
        'first_name, last_name - Individual names',
        'organization_name - Company/organization',
        'role - Role in application (owner, designer, agent)',
        'address, city, province_state, postal_code - Contact information',
        'country, country_code - Country information'
    ]
    for field in fields:
        doc.add_paragraph(field, style='List Bullet')

    add_heading(doc, 'Other Application Models', 3)
    add_paragraph(doc, 'The following models provide additional application details:')
    models = [
        'ApplicationImage - Image assets associated with applications',
        'ApplicationDescription - Textual descriptions of industrial designs',
        'ApplicationDescriptionTxtFormat - Alternative text-formatted descriptions',
        'ApplicationCorrection - Records corrections made to published applications'
    ]
    for model in models:
        doc.add_paragraph(model, style='List Bullet')

    add_heading(doc, '4.2 Assignment Models', 2)

    add_heading(doc, 'AssignmentMain', 3)
    add_paragraph(doc, 'Records of ownership transfers and assignments.')
    add_paragraph(doc, 'Key Fields:')
    fields = [
        'assignment_number - Unique assignment identifier',
        'application_number - Links to ApplicationMain',
        'assignment_type, assignment_type_code - Type of assignment',
        'assignment_status - Current status',
        'assignment_registration_date - When recorded',
        'ownership_change_prior_to_filing - Pre-filing ownership changes'
    ]
    for field in fields:
        doc.add_paragraph(field, style='List Bullet')

    add_heading(doc, 'Other Assignment Models', 3)
    models = [
        'AssignmentInterestedParty - Parties involved in assignments',
        'AssignmentCorrection - Corrections to published assignments'
    ]
    for model in models:
        doc.add_paragraph(model, style='List Bullet')

    # 5. API Endpoints
    add_heading(doc, '5. API Endpoints', 1)
    add_paragraph(doc,
        'All endpoints are prefixed with /api/ and support standard REST operations (GET only).')

    add_heading(doc, '5.1 Base Endpoints', 2)
    table_data = [
        '| Endpoint | Purpose | Model |',
        '|----------|---------|-------|',
        '| /api/main/ | Core application data | ApplicationMain |',
        '| /api/application_classification/ | Product classifications | ApplicationClassification |',
        '| /api/application_correction/ | Application corrections | ApplicationCorrection |',
        '| /api/application_description/ | Design descriptions | ApplicationDescription |',
        '| /api/application_image/ | Image metadata | ApplicationImage |',
        '| /api/application_interested_party/ | Application parties | ApplicationInterestedParty |',
        '| /api/assignment_main/ | Assignment records | AssignmentMain |',
        '| /api/assignment_interested_party/ | Assignment parties | AssignmentInterestedParty |',
        '| /api/assignment_correction/ | Assignment corrections | AssignmentCorrection |'
    ]
    add_table_from_markdown(doc, table_data)

    add_heading(doc, '5.2 Query Parameters', 2)

    add_heading(doc, 'Filtering', 3)
    add_paragraph(doc, 'Field-based filters:')
    add_code_block(doc, '?field_name=value')

    add_paragraph(doc, 'Date range filters:')
    add_code_block(doc, '?filing_date_after=2023-01-01\n?filing_date_before=2024-12-31')

    add_heading(doc, 'Search', 3)
    add_code_block(doc, '?search=keyword')
    add_paragraph(doc, 'Searches across all text fields defined in the viewset.')

    add_heading(doc, 'Ordering', 3)
    add_code_block(doc, '?ordering=field_name\n?ordering=-field_name  # Descending')

    add_heading(doc, 'Omit Filter', 3)
    add_code_block(doc, '?omit=keyword\n?omit=field:keyword\n?omit=word1,word2')
    add_paragraph(doc, 'Excludes records containing specified keywords.')

    add_heading(doc, '5.3 API Documentation', 2)
    add_paragraph(doc, 'Interactive documentation available at:')
    docs = [
        'Swagger UI: /swagger/',
        'ReDoc: /redoc/',
        'OpenAPI JSON: /swagger.json',
        'OpenAPI YAML: /swagger.yaml'
    ]
    for doc_item in docs:
        doc.add_paragraph(doc_item, style='List Bullet')

    # 6. API Usage Examples
    add_heading(doc, '6. API Usage Examples', 1)

    add_heading(doc, '6.1 Basic Queries', 2)

    add_paragraph(doc, 'Get all applications:')
    add_code_block(doc, 'curl http://api.opic-cipo.ca/api/main/')

    add_paragraph(doc, 'Get specific application by ID:')
    add_code_block(doc, 'curl http://api.opic-cipo.ca/api/main/12345/')

    add_paragraph(doc, 'Get application by application number:')
    add_code_block(doc, 'curl http://api.opic-cipo.ca/api/main/by-number/123456/')

    add_heading(doc, '6.2 Search and Filter', 2)

    add_paragraph(doc, 'Search for applications:')
    add_code_block(doc, 'curl "http://api.opic-cipo.ca/api/main/?search=furniture"')

    add_paragraph(doc, 'Filter by date range:')
    add_code_block(doc,
        'curl "http://api.opic-cipo.ca/api/main/\n     ?filing_date_after=2024-01-01\n     &filing_date_before=2024-12-31"')

    add_paragraph(doc, 'Filter and sort:')
    add_code_block(doc,
        'curl "http://api.opic-cipo.ca/api/main/\n     ?design_status=Registered\n     &ordering=-filing_date"')

    # 7. Installation and Setup
    add_heading(doc, '7. Installation and Setup', 1)

    add_heading(doc, '7.1 Prerequisites', 2)
    prereqs = [
        'Python 3.8+',
        'PostgreSQL 12+',
        'pip package manager',
        'Virtual environment (recommended)'
    ]
    for prereq in prereqs:
        doc.add_paragraph(prereq, style='List Bullet')

    add_heading(doc, '7.2 Installation Steps', 2)

    add_paragraph(doc, '1. Clone the repository:')
    add_code_block(doc, 'git clone <repository-url>\ncd dd_industrial_design_API')

    add_paragraph(doc, '2. Create virtual environment:')
    add_code_block(doc,
        'python -m venv venv\nsource venv/bin/activate  # On Windows: venv\\Scripts\\activate')

    add_paragraph(doc, '3. Install dependencies:')
    add_code_block(doc, 'pip install -r requirements.txt')

    add_paragraph(doc, '4. Configure database in DB_Main/settings.py')

    add_paragraph(doc, '5. Test the configuration:')
    add_code_block(doc, 'python manage.py check')

    add_paragraph(doc, '6. Run development server:')
    add_code_block(doc, 'python manage.py runserver')

    add_paragraph(doc, 'The API will be available at http://localhost:8000/')

    # 8. Configuration
    add_heading(doc, '8. Configuration', 1)

    add_heading(doc, '8.1 Key Settings', 2)

    add_paragraph(doc, 'Debug Mode (settings.py:169):')
    add_code_block(doc, "DEBUG = True  # Set to False in production")

    add_paragraph(doc, 'Allowed Hosts (settings.py:223):')
    add_code_block(doc, "ALLOWED_HOSTS = ['127.0.0.1', 'localhost', 'api.opic-cipo.ca']")

    add_paragraph(doc, 'Pagination (settings.py:174-175):')
    add_code_block(doc,
        "REST_FRAMEWORK = {\n    'DEFAULT_PAGINATION_CLASS': 'rest_framework.pagination.PageNumberPagination',\n    'PAGE_SIZE': 50,\n}")

    # 9. Deployment
    add_heading(doc, '9. Deployment', 1)

    add_heading(doc, '9.1 Production Server', 2)
    add_paragraph(doc,
        'The application uses Waitress as the production WSGI server.')

    add_paragraph(doc, 'Start production server:')
    add_code_block(doc, 'python serve.py')

    add_paragraph(doc, 'Server configuration (serve.py):')
    config = [
        'Host: 0.0.0.0 (all interfaces)',
        'Port: 8080'
    ]
    for item in config:
        doc.add_paragraph(item, style='List Bullet')

    add_heading(doc, '9.2 Production Checklist', 2)
    checklist = [
        'Set DEBUG = False in settings.py',
        'Configure appropriate ALLOWED_HOSTS',
        'Set secure SECRET_KEY',
        'Configure static file serving',
        'Set up SSL/TLS certificates',
        'Configure firewall rules',
        'Set up database backups',
        'Configure logging and monitoring',
        'Set appropriate file permissions',
        'Use environment variables for sensitive data'
    ]
    for item in checklist:
        doc.add_paragraph(item, style='List Number')

    # 10. Code Architecture
    add_heading(doc, '10. Code Architecture', 1)

    add_heading(doc, '10.1 Custom Filter Backend', 2)
    add_paragraph(doc, 'OmitFilterBackend (views.py:35-66)')
    add_paragraph(doc,
        'Custom Django filter that excludes records containing specified keywords.')

    add_heading(doc, '10.2 Dynamic FilterSet Generation', 2)
    add_paragraph(doc, 'make_ab_filterset() (views.py:75-89)')
    add_paragraph(doc,
        'Dynamically generates Django FilterSet classes with _after and _before range filters '
        'for date and numeric fields.')

    add_heading(doc, '10.3 ReadOnlyMixin', 2)
    add_paragraph(doc, 'ReadOnlyMixin (views.py:91-113)')
    add_paragraph(doc, 'Base viewset providing:')
    features = [
        'Read-only operations (list, retrieve)',
        'Dynamic filterset generation',
        'Ordering and search capabilities',
        'Custom by-number lookup endpoint'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')

    # 11. Security Considerations
    add_heading(doc, '11. Security Considerations', 1)

    add_heading(doc, '11.1 Current Configuration', 2)
    security = [
        'Authentication: None (public read-only API)',
        'CORS: Configured for specific domains',
        'SQL Injection: Protected by Django ORM',
        'XSS: JSON-only responses (no HTML rendering)'
    ]
    for item in security:
        doc.add_paragraph(item, style='List Bullet')

    add_heading(doc, '11.2 Recommendations', 2)
    recommendations = [
        'Implement rate limiting for production',
        'Add API key authentication for tracking usage',
        'Enable HTTPS only',
        'Implement request logging',
        'Set up monitoring and alerting',
        'Regular security updates'
    ]
    for item in recommendations:
        doc.add_paragraph(item, style='List Number')

    # 12. Troubleshooting
    add_heading(doc, '12. Troubleshooting', 1)

    add_heading(doc, '12.1 Common Issues', 2)

    add_paragraph(doc, 'Database Connection Errors:')
    solutions = [
        'Verify PostgreSQL is running',
        'Check credentials in settings.py',
        'Ensure database exists: IndustrialDesign',
        'Verify user permissions'
    ]
    for solution in solutions:
        doc.add_paragraph(solution, style='List Bullet')

    add_paragraph(doc, 'Import Errors:')
    solutions = [
        'Activate virtual environment',
        'Reinstall requirements: pip install -r requirements.txt'
    ]
    for solution in solutions:
        doc.add_paragraph(solution, style='List Bullet')

    add_paragraph(doc, '404 on API Endpoints:')
    solutions = [
        'Verify URL includes /api/ prefix',
        'Check swagger documentation for exact paths'
    ]
    for solution in solutions:
        doc.add_paragraph(solution, style='List Bullet')

    # 13. Knowledge Transfer Notes
    add_heading(doc, '13. Knowledge Transfer Notes', 1)

    add_heading(doc, '13.1 Key Understanding Points', 2)

    points = [
        'Read-Only Nature: API provides only GET operations. No data modification capabilities.',
        'Unmanaged Models: Database schema is controlled externally, not by Django migrations.',
        'Schema Organization: All data resides in PostgreSQL schema id_csv_2024_03_07.',
        'Filtering Architecture: Three-tier filtering system (Django Filter, Search, Omit).',
        'Number Lookups: Special endpoint pattern /by-number/{number}/ for queries.'
    ]
    for point in points:
        doc.add_paragraph(point, style='List Number')

    add_heading(doc, '13.2 Development Workflow', 2)

    add_paragraph(doc, 'Making Changes: Modify views, serializers, or add endpoints in industrial_designs app')
    add_paragraph(doc, 'Testing: Use Swagger UI at /swagger/ for interactive testing')
    add_paragraph(doc, 'Deployment: Run python serve.py for production deployment')

    add_heading(doc, '13.3 Extension Points', 2)
    add_paragraph(doc, 'To add new endpoints:')
    steps = [
        'Define model in models.py (set managed = False)',
        'Create serializer in serializers.py',
        'Create viewset in views.py (inherit from ReadOnlyMixin)',
        'Register in urls.py router'
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')

    # Document Information
    doc.add_page_break()
    add_heading(doc, 'Document Information', 1)
    doc.add_paragraph('Document Version: 1.0')
    doc.add_paragraph('Last Updated: 2025-11-25')
    doc.add_paragraph('Project Repository: dd_industrial_design_API')
    doc.add_paragraph('Database Schema Version: id_csv_2024_03_07')

    # Save document
    output_file = '/home/user/dd_industrial_design_API/Industrial_Design_API_Technical_Documentation.docx'
    doc.save(output_file)
    print(f"Word document created successfully: {output_file}")
    return output_file

if __name__ == '__main__':
    generate_word_document()
